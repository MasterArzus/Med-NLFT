import json
import os
from docx import Document
import re
import argparse

# 提取加粗的关键词
def extract_bold_text(doc):
    bold_texts = []
    for para in doc.paragraphs:
        bold_run = ""
        
        for run in para.runs:
            if run.bold:
                bold_run += run.text  # 拼接加粗的文字
            else:
                if bold_run:  # 如果当前有加粗的文本被拼接
                    bold_texts.append(bold_run)  # 保存拼接的加粗文本
                    bold_run = ""  # 清空已拼接的内容

        # 如果段落结束时还有未保存的加粗文本
        if bold_run:
            bold_texts.append(bold_run)

    return bold_texts

# 提取两个关键词间的文字
def extract_text_between(doc, keyword1, keyword2):
    content = ""
    found_keyword1 = False

    # 遍历文档中的每个段落
    for para in doc.paragraphs:
        if keyword1 in para.text and not found_keyword1:
            # 找到第一个关键词后标记
            found_keyword1 = True
            content = para.text.split(keyword1, 1)[-1]  # 从第一个关键词开始获取内容
        elif found_keyword1:
            # 如果已经找到了第一个关键词，继续收集内容直到找到第二个关键词
            content += '\n' + para.text
            if keyword2 in para.text:
                content = content.split(keyword2, 1)[0]  # 获取到第二个关键词时停止
                break

    # 输出找到的内容
    if content:
        return content
    else:
        return "未找到这两个关键词之间的内容"

# 将文本中的中文标点符号转换成英文
def clean_punctuation(text):
    # 定义中文符号与英文符号的映射
    punctuation_map = {
        '，': ',', '。': '.', '！': '!', '？': '?', '；': ';', '：': ':',
        '“': '"', '”': '"', '‘': "'", '’': "'", '【': '[', '】': ']',
        '（': '(', '）': ')', '《': '<', '》': '>', '、': ',', '——': '-',
        '／': '/',
    }

    # 使用正则表达式匹配中文符号并进行替换
    pattern = re.compile('|'.join(re.escape(key) for key in punctuation_map.keys()))
    
    def replace(match):
        return punctuation_map[match.group(0)]

    return pattern.sub(replace, text)

# 去除首尾的符号和空格,删除内部多余的空格
def clean_string(s):
    s = s.strip(' \t\n\r\f\v`~!@#$%^&*()_+-=[]{}|;:\'",.<>?/\\')
    s = re.sub(r'\s+', ' ', s)
    return s

def extract_data_from_docx(doc_path):
    doc = Document(doc_path)
    data = []
    ques_key=[]
    ans_key=['诊断','用药情况','是否根据病原学结果调整','出院后是否继续治疗','治疗结局','住院天数','出院带药情况']

    question = {}
    answer = {}
    bold_texts = extract_bold_text(doc)
    bold_texts.append('\n')
    # print(bold_texts)

    for i in range(len(bold_texts) - 1):
        keyword1 = bold_texts[i]
        keyword2 = bold_texts[i + 1]
        key = clean_string(clean_punctuation(keyword1))
        content = extract_text_between(doc, keyword1, keyword2)
        value = clean_string(clean_punctuation(content))

        if key == '细菌培养及药敏':
            for table in doc.tables:
                # 遍历表格中的每一行
                for row in table.rows:
                    if '>' in row.cells[1].text:
                        value += ' '+row.cells[0].text+row.cells[1].text+' '
                        value = clean_punctuation(value)
                        # print(row.cells[0].text, row.cells[1].text)

        if len(content) > 1:
            if key not in ans_key:
                question.update({key:value})
            else:
                answer.update({key:value})
        else:
            continue


    # 合并为一个字典
    data.append({"question": question, "answer": answer})
    
    return data


def save_to_json(new_data, filepath):
    # 检查文件是否存在并读取数据
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                try:
                    existing_data = json.load(f)
                except json.JSONDecodeError:
                    existing_data = []
        except FileNotFoundError:
            existing_data = []
    else:
        existing_data = []
    
    # 确保数据为列表，并追加新数据
    if not isinstance(existing_data, list):
        existing_data = [existing_data]
    existing_data.append(new_data)
    
    # 写回文件
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(existing_data, f, indent=4, ensure_ascii=False)
    return True


def main():

    parser = argparse.ArgumentParser(description='create dataset for Med-Info')
    parser.add_argument('--input_word', type=str, required=True, help='path to input word')
    parser.add_argument('--dataset_dir', type=str, required=True, help='path to dataset dir')
    parser.add_argument('--n_total', type=int, default=50, help='total number of data') 
    parser.add_argument('--ratio', type=str, default=None, choices=['correct', '1:3', '1:2', '1:1', '2:1', '3:1', 'wrong', None], 
                        help='ratio of correct and wrong') # 对比错的数目比例

    args = parser.parse_args()
    input_file_path = args.input_word
    dataset_dir = args.dataset_dir
    n_total = args.n_total
    ratio = args.ratio

    med_dir = os.path.join(dataset_dir, 'Med')
    os.makedirs(med_dir, exist_ok=True)

    med_filename = "{}Med{}.json".format(ratio.replace(':', '_') + '_' if ratio else '', n_total)
    med_path = os.path.join(med_dir, med_filename)

    data = extract_data_from_docx(input_file_path)
    save_to_json(data, med_path)

    print("{} processed.".format(med_filename))


if __name__ == '__main__':
    main()
